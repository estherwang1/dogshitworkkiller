"""
word 文档遍历器

按 body 顺序遍历 word 文档,返回块级元素列表。
段落和表格按真实交错顺序编号,为两条链路提供统一的数据基础:
- 链路一(评估):拼接 text_full,生成供评估模型阅读的纯文本
- 链路二(分块标注):格式化为带编号的文本,供模型识别节标题位置

依赖:
    pip install python-docx

使用:
    from word_parser import parse_docx, format_for_annotation, join_for_evaluation

    blocks = parse_docx("某标准.docx")
    # 链路二:
    text_for_model = format_for_annotation(blocks)
    # 链路一:
    plain_text = join_for_evaluation(blocks)
"""
from typing import Optional
from docx import Document
from docx.oxml.ns import qn


# XML 命名空间常量
W_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TAG_PARAGRAPH = f"{{{W_NAMESPACE}}}p"
TAG_TABLE = f"{{{W_NAMESPACE}}}tbl"

# 样式三分类
# 强正信号:极可能是节/章标题
HEADING_STYLE_PREFIXES = (
    "heading",        # Heading 1-9
    "title",          # Title / Subtitle
    "标题",
    "副标题",
    "章标题", "节标题", "条标题",
    "一级标题", "二级标题", "三级标题",
)

# 强负信号:明确不是正文标题
NON_HEADING_STYLE_PREFIXES = (
    "toc",
    "caption",
    "目录",
    "图注", "表注",
)


def _classify_style(style_name: str) -> str:
    """把样式名归类为三种之一

    返回:
        "heading"     - 标题样式(强正信号)
        "non_heading" - 目录/图注类(强负信号)
        "other"       - 正文或其他(弱信号,不输出给模型)
    """
    if not style_name:
        return "other"
    s = style_name.strip().lower()
    for prefix in HEADING_STYLE_PREFIXES:
        if s.startswith(prefix.lower()):
            return "heading"
    for prefix in NON_HEADING_STYLE_PREFIXES:
        if s.startswith(prefix.lower()):
            return "non_heading"
    return "other"


def _has_inline_image(paragraph) -> bool:
    """检测段落是否含 inline 图片

    只识别标准的 drawing 元素,不处理文本框、SmartArt、VML 形状。
    """
    for run in paragraph.runs:
        if run.element.findall(qn("w:drawing")):
            return True
    return False


# DrawingML 命名空间(用于 SmartArt、形状等)
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _extract_shape_texts(paragraph) -> list[str]:
    """提取段落内图形对象(文本框、形状、SmartArt 等)里的文字

    涵盖范围:
    - 文本框:<w:txbxContent> 下的 <w:t>
    - 形状组合、画布、SmartArt 里的文本:<a:t> 标签(DrawingML 通用文本)
    - 老版本 VML 文本框:<w:txbxContent> 通常也能覆盖

    返回去重后的非空文本列表,保留原始顺序。
    不区分具体类型,只负责"把能提取的文字捞出来"。
    """
    element = paragraph._element
    texts = []
    seen = set()

    # 1. word 文本框(较常见)
    for t in element.iter(qn("w:txbxContent")):
        for text_elem in t.iter(qn("w:t")):
            s = (text_elem.text or "").strip()
            if s and s not in seen:
                seen.add(s)
                texts.append(s)

    # 2. DrawingML 通用文本(SmartArt、形状里的文字)
    for text_elem in element.iter(f"{{{A_NS}}}t"):
        s = (text_elem.text or "").strip()
        if s and s not in seen:
            seen.add(s)
            texts.append(s)

    return texts


def _has_graphic_object(paragraph) -> bool:
    """检测段落是否含非 inline 图片的图形对象(文本框/形状/SmartArt/画布等)

    判断依据:
    - 存在 <mc:AlternateContent>(wps/wpg/wpc 类形状)
    - 存在 <w:txbxContent>(文本框)
    - 存在 DrawingML 的 diagram/chart(SmartArt/图表)

    注:如果段落里只有 inline 图片(<w:drawing>),不算这里的"图形对象",
    inline 图片由 _has_inline_image 单独识别。
    """
    element = paragraph._element
    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

    # AlternateContent 下常见形状/画布/组合
    if list(element.iter(f"{{{MC_NS}}}AlternateContent")):
        return True
    # 独立文本框
    if list(element.iter(qn("w:txbxContent"))):
        return True
    # DrawingML diagram(SmartArt)
    for gd in element.iter(f"{{{A_NS}}}graphicData"):
        uri = gd.get("uri") or ""
        if "diagram" in uri or "chart" in uri:
            return True
    return False


def _is_all_runs_bold(paragraph) -> bool:
    """判断段落是否整段加粗(所有有文字的 run 都显式加粗)"""
    runs_with_text = [r for r in paragraph.runs if r.text]
    if not runs_with_text:
        return False
    return all(r.bold is True for r in runs_with_text)


def _get_first_font_size(paragraph) -> Optional[float]:
    """读段落第一个显式设置字号的 run 的字号(pt)

    不做样式继承链递归,读不到就返回 None。
    """
    for run in paragraph.runs:
        if run.font.size is not None:
            return run.font.size.pt
    return None


def _get_style_name(paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def _paragraph_text(paragraph) -> str:
    return paragraph.text or ""


def _format_graphic_marker(texts: list[str]) -> str:
    """把图形对象的文字列表格式化成标记字符串

    示例:
    - []             → "<含图形对象>"
    - ["图 1"]        → '<含图形对象: "图 1">'
    - ["A", "B", "C"] → '<含图形对象: "A", "B", "C">'
    """
    if not texts:
        return "<含图形对象>"
    quoted = [f'"{t}"' for t in texts]
    return f"<含图形对象: {', '.join(quoted)}>"


def _classify_paragraph(paragraph) -> tuple[str, str]:
    """判断段落类型,返回 (type, display_text)

    判断维度:
    - 文字(w:t 里的正文)
    - inline 图片(w:drawing)
    - 图形对象(文本框/形状/SmartArt 等,单独检测)

    返回规则:
    - 无文字、无 inline 图片、无图形对象  → empty, "<空段>"
    - 无文字、有 inline 图片、无图形对象  → image, "<图片>"
    - 其他情况均为 paragraph,显示文本按检测到的元素追加标记

    图形对象的文字会提取出来附在标记里,例如:
      正文xxx <含图形对象: "图 1", "组件 A">
    """
    text = _paragraph_text(paragraph).strip()
    has_image = _has_inline_image(paragraph)
    has_graphic = _has_graphic_object(paragraph)
    graphic_texts = _extract_shape_texts(paragraph) if has_graphic else []

    # 纯空
    if not text and not has_image and not has_graphic:
        return "empty", "<空段>"

    # 纯 inline 图片(无文字无图形对象)
    if not text and has_image and not has_graphic:
        return "image", "<图片>"

    # 其余都是 paragraph,按元素组合拼显示文本
    parts = []
    if text:
        parts.append(text)
    if has_image:
        parts.append("<含图片>")
    if has_graphic:
        parts.append(_format_graphic_marker(graphic_texts))

    return "paragraph", " ".join(parts)


def _table_to_markdown(table) -> str:
    """把表格转成 markdown 格式字符串

    合并单元格不去重,保留原始返回。
    单元格内容里的 | 替换成 /。
    """
    rows_text = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = " ".join(p.text for p in cell.paragraphs if p.text)
            cell_text = cell_text.replace("|", "/")
            cells.append(cell_text)
        rows_text.append("| " + " | ".join(cells) + " |")

    if not rows_text:
        return "<空表格>"

    col_count = len(table.rows[0].cells) if table.rows else 0
    if col_count > 0 and len(rows_text) >= 1:
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows_text.insert(1, separator)

    return "\n".join(rows_text)


def parse_docx(path: str) -> list[dict]:
    """遍历 word 文档,返回块级元素列表

    每个元素字段:
    - id: "P0001" 补零 4 位,按 body 顺序连续编号
    - type: "paragraph" / "table" / "image" / "empty"
    - text: 显示文本,含 <含图片>/<图片>/<空段>/<表格: NxM> 等标记
    - text_full: 完整文本,仅表格与 text 不同
    - style: 段落样式名,表格和非段落为 ""
    - style_class: "heading" / "non_heading" / "other"
    - is_bold: 段落是否整段加粗
    - font_size: 段落第一个显式字号(pt)
    """
    doc = Document(path)
    body = doc.element.body
    blocks = []
    index = 0

    paragraph_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag == TAG_PARAGRAPH:
            paragraph = paragraph_map.get(child)
            if paragraph is None:
                continue
            index += 1
            block_type, display_text = _classify_paragraph(paragraph)
            style_name = _get_style_name(paragraph)
            blocks.append({
                "id": f"P{index:04d}",
                "type": block_type,
                "text": display_text,
                "text_full": display_text,
                "style": style_name,
                "style_class": _classify_style(style_name),
                "is_bold": _is_all_runs_bold(paragraph),
                "font_size": _get_first_font_size(paragraph),
            })
        elif child.tag == TAG_TABLE:
            table = table_map.get(child)
            if table is None:
                continue
            index += 1
            row_count = len(table.rows)
            col_count = len(table.rows[0].cells) if row_count > 0 else 0
            blocks.append({
                "id": f"P{index:04d}",
                "type": "table",
                "text": f"<表格: {row_count}行{col_count}列>",
                "text_full": _table_to_markdown(table),
                "style": "",
                "style_class": "other",
                "is_bold": False,
                "font_size": None,
            })

    return blocks


def _format_font_size(size: Optional[float]) -> str:
    """字号显示:整数去掉 .0"""
    if size is None:
        return ""
    if size == int(size):
        return str(int(size))
    return str(size)


def _format_paragraph_signals(block: dict) -> str:
    """格式化段落的格式信号部分

    策略:
    - style_class == "heading"     → 输出 "style=Heading 2"
    - style_class == "non_heading" → 输出 "style=TOC 2 (目录/图注)"
    - style_class == "other"       → 不输出 style(降噪)
    - 加粗和字号信息在 heading 和 other 下都输出(作为伪标题检测依据)
    - non_heading 不输出加粗/字号(反正要排除)
    """
    signals = []
    style_class = block["style_class"]

    if style_class == "heading":
        signals.append(f"style={block['style']}")
    elif style_class == "non_heading":
        signals.append(f"style={block['style']} (目录/图注)")

    if style_class != "non_heading":
        if block["is_bold"]:
            signals.append("加粗")
        if block["font_size"] is not None:
            signals.append(f"字号={_format_font_size(block['font_size'])}")

    return ", ".join(signals)


def format_for_annotation(blocks: list[dict]) -> str:
    """把块列表格式化为带编号的文本,供链路二的识别 prompt 使用

    格式:
        [P0001] 3.2 密码算法要求 | style=Heading 2
        [P0002] 正文内容
        [P0003-P0012] <连续10个空段>
        [P0013] 下一段

    折叠规则:
    - 连续 >= 2 个空段折叠为 [P0003-P0012] <连续N个空段>
    - 单个空段保留原样
    - 表格/图片不折叠
    """
    lines = []
    i = 0
    n = len(blocks)

    while i < n:
        block = blocks[i]

        if block["type"] == "empty":
            j = i
            while j < n and blocks[j]["type"] == "empty":
                j += 1
            run_length = j - i

            if run_length >= 2:
                first_id = blocks[i]["id"]
                last_id = blocks[j - 1]["id"]
                lines.append(f"[{first_id}-{last_id}] <连续{run_length}个空段>")
                i = j
                continue
            else:
                lines.append(f"[{block['id']}] <空段>")
                i += 1
                continue

        if block["type"] in ("table", "image"):
            lines.append(f"[{block['id']}] {block['text']}")
        else:
            signals = _format_paragraph_signals(block)
            if signals:
                lines.append(f"[{block['id']}] {block['text']} | {signals}")
            else:
                lines.append(f"[{block['id']}] {block['text']}")
        i += 1

    return "\n".join(lines)


def join_for_evaluation(blocks: list[dict]) -> str:
    """把块列表拼成纯文本,供链路一的评估 prompt 使用

    不含编号,不含格式信号。
    """
    parts = [block["text_full"] for block in blocks]
    return "\n\n".join(parts)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python word_parser.py 某标准.docx")
        sys.exit(1)

    blocks = parse_docx(sys.argv[1])
    print(f"共提取 {len(blocks)} 个块级元素\n")

    from collections import Counter
    type_count = Counter(b["type"] for b in blocks)
    style_class_count = Counter(
        b["style_class"] for b in blocks if b["type"] == "paragraph"
    )
    print("块类型分布:")
    for t, c in type_count.most_common():
        print(f"  {t:12s} {c}")
    print("\n样式分类分布(仅段落):")
    for sc, c in style_class_count.most_common():
        print(f"  {sc:12s} {c}")

    print("\n--- format_for_annotation 前 30 行 ---")
    formatted = format_for_annotation(blocks)
    print("\n".join(formatted.split("\n")[:30]))
